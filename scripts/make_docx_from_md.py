#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path


def add_md_to_doc(doc, md: str) -> None:
    """
    Minimal Markdown-to-DOCX conversion for this project's simple structure:
    - '#', '##', '###' headings
    - bullet lists starting with '- '
    - blank lines separate paragraphs
    """
    in_list = False
    for raw in md.splitlines():
        line = raw.rstrip("\n")
        if not line.strip():
            in_list = False
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            in_list = False
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            in_list = False
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            in_list = False
            continue

        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            in_list = True
            continue

        # treat everything else as normal paragraph
        doc.add_paragraph(line.strip())
        in_list = False


def main() -> int:
    repo_root = Path(__file__).resolve().parent.parent
    parser = argparse.ArgumentParser(description="Convert a simple markdown file to a .docx.")
    parser.add_argument(
        "--md",
        type=str,
        default=str(repo_root / "EVALUATION_NEXT_STEPS.md"),
        help="Input markdown path",
    )
    parser.add_argument(
        "--out",
        type=str,
        default=str(repo_root / "EVALUATION_NEXT_STEPS.docx"),
        help="Output docx path",
    )
    args = parser.parse_args()

    md_path = Path(args.md)
    out_path = Path(args.out)

    if not md_path.exists():
        raise FileNotFoundError(f"Missing input: {md_path}")

    from docx import Document

    md = md_path.read_text(encoding="utf-8", errors="replace")
    doc = Document()
    add_md_to_doc(doc, md)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

